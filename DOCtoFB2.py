#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
DOCtoFB2 - Конвертер DOC/DOCX в FB2 для Литрес Самиздат
Версия 1.0
"""

import sys
import os
import json
import zipfile
import base64
import re
import tempfile
from pathlib import Path
from datetime import datetime
from typing import Optional, List, Dict, Any
from dataclasses import dataclass, asdict
from io import BytesIO

# Проверка зависимостей
try:
    from docx import Document
    from docx.enum.style import WD_STYLE_TYPE
    from docx.shared import Inches
    from lxml import etree
    from PIL import Image
    from PyQt5.QtWidgets import (
        QApplication, QMainWindow, QWidget, QVBoxLayout, QHBoxLayout,
        QPushButton, QLabel, QFileDialog, QTextEdit, QSplitter,
        QToolBar, QAction, QStatusBar, QMessageBox, QDialog,
        QListWidget, QListWidgetItem, QGroupBox, QCheckBox,
        QLineEdit, QFormLayout, QDialogButtonBox, QTabWidget
    )
    from PyQt5.QtCore import Qt, QSettings, QSize, QMimeData, QUrl
    from PyQt5.QtGui import QIcon, QFont, QDragEnterEvent, QDropEvent
except ImportError as e:
    print("Установите необходимые зависимости:")
    print("pip install python-docx lxml pillow pyqt5")
    print(f"Ошибка: {e}")
    sys.exit(1)


@dataclass
class AppSettings:
    """Настройки приложения"""
    default_save_path: str = ""
    preserve_formatting: bool = True
    convert_images: bool = True
    remove_empty_lines: bool = True
    auto_open_result: bool = False
    validate_fb2: bool = True
    
    @classmethod
    def load(cls) -> 'AppSettings':
        """Загрузить настройки из файла"""
        settings_file = Path.home() / '.doctofb2_settings.json'
        if settings_file.exists():
            try:
                with open(settings_file, 'r', encoding='utf-8') as f:
                    data = json.load(f)
                    return cls(**data)
            except:
                pass
        return cls()
    
    def save(self):
        """Сохранить настройки в файл"""
        settings_file = Path.home() / '.doctofb2_settings.json'
        with open(settings_file, 'w', encoding='utf-8') as f:
            json.dump(asdict(self), f, ensure_ascii=False, indent=2)


class SettingsDialog(QDialog):
    """Диалог настроек"""
    
    def __init__(self, parent=None):
        super().__init__(parent)
        self.settings = AppSettings.load()
        self.init_ui()
        
    def init_ui(self):
        self.setWindowTitle("Настройки")
        self.setFixedSize(500, 400)
        
        layout = QVBoxLayout()
        
        # Вкладки
        tabs = QTabWidget()
        
        # Основные настройки
        basic_tab = QWidget()
        basic_layout = QFormLayout()
        
        self.save_path_edit = QLineEdit(self.settings.default_save_path)
        self.save_path_edit.setPlaceholderText("Путь для сохранения по умолчанию")
        browse_btn = QPushButton("Обзор...")
        browse_btn.clicked.connect(self.browse_save_path)
        
        path_layout = QHBoxLayout()
        path_layout.addWidget(self.save_path_edit)
        path_layout.addWidget(browse_btn)
        
        basic_layout.addRow("Путь сохранения:", path_layout)
        
        self.preserve_cb = QCheckBox("Сохранять форматирование")
        self.preserve_cb.setChecked(self.settings.preserve_formatting)
        basic_layout.addRow(self.preserve_cb)
        
        self.convert_images_cb = QCheckBox("Конвертировать изображения в JPG")
        self.convert_images_cb.setChecked(self.settings.convert_images)
        basic_layout.addRow(self.convert_images_cb)
        
        self.remove_empty_cb = QCheckBox("Удалять пустые строки")
        self.remove_empty_cb.setChecked(self.settings.remove_empty_lines)
        basic_layout.addRow(self.remove_empty_cb)
        
        self.auto_open_cb = QCheckBox("Автоматически открывать результат")
        self.auto_open_cb.setChecked(self.settings.auto_open_result)
        basic_layout.addRow(self.auto_open_cb)
        
        self.validate_cb = QCheckBox("Проверять валидность FB2")
        self.validate_cb.setChecked(self.settings.validate_fb2)
        basic_layout.addRow(self.validate_cb)
        
        basic_tab.setLayout(basic_layout)
        tabs.addTab(basic_tab, "Основные")
        
        # Правила Литрес
        rules_tab = QWidget()
        rules_layout = QVBoxLayout()
        
        rules_text = QTextEdit()
        rules_text.setReadOnly(True)
        rules_text.setHtml("""
        <h3>Правила подготовки файлов для Литрес:</h3>
        <ol>
        <li><b>Заголовки:</b> Используйте стили "Заголовок 1", "Заголовок 2", "Заголовок 3"</li>
        <li><b>Пустые строки:</b> Не оставляйте пустых строк после заголовков</li>
        <li><b>Изображения:</b> Вставляйте через "Вставка → Рисунки → Из файла"</li>
        <li><b>Формат обтекания:</b> Устанавливайте "В тексте"</li>
        <li><b>Сноски:</b> Используйте функцию "Сноска" в Word</li>
        <li><b>Таблицы:</b> Конвертируйте в изображения</li>
        <li><b>Эмодзи:</b> Не используйте, заменяйте на текстовые описания</li>
        <li><b>Символы:</b> Избегайте специальных символов и символов иностранных языков в заголовках</li>
        </ol>
        <p><i>Для заголовков с иностранными символами наберите текст в Блокноте, 
        затем скопируйте в Word и примените стиль заголовка.</i></p>
        """)
        rules_layout.addWidget(rules_text)
        rules_tab.setLayout(rules_layout)
        tabs.addTab(rules_tab, "Правила Литрес")
        
        layout.addWidget(tabs)
        
        # Кнопки
        buttons = QDialogButtonBox(
            QDialogButtonBox.Ok | QDialogButtonBox.Cancel | QDialogButtonBox.Apply
        )
        buttons.accepted.connect(self.accept)
        buttons.rejected.connect(self.reject)
        buttons.button(QDialogButtonBox.Apply).clicked.connect(self.apply_settings)
        
        layout.addWidget(buttons)
        self.setLayout(layout)
    
    def browse_save_path(self):
        """Выбор пути для сохранения"""
        path = QFileDialog.getExistingDirectory(self, "Выберите папку для сохранения")
        if path:
            self.save_path_edit.setText(path)
    
    def apply_settings(self):
        """Применить настройки"""
        self.settings.default_save_path = self.save_path_edit.text()
        self.settings.preserve_formatting = self.preserve_cb.isChecked()
        self.settings.convert_images = self.convert_images_cb.isChecked()
        self.settings.remove_empty_lines = self.remove_empty_cb.isChecked()
        self.settings.auto_open_result = self.auto_open_cb.isChecked()
        self.settings.validate_fb2 = self.validate_cb.isChecked()
        self.settings.save()
        QMessageBox.information(self, "Настройки", "Настройки сохранены!")
    
    def accept(self):
        """Принять изменения"""
        self.apply_settings()
        super().accept()


class DocxToFb2Converter:
    """Класс для конвертации DOCX в FB2"""
    
    def __init__(self):
        self.images = []
        self.image_counter = 0
    
    def convert(self, docx_path: str, settings: AppSettings) -> str:
        """Конвертировать DOCX в FB2"""
        try:
            # Чтение DOCX
            doc = Document(docx_path)
            
            # Создание структуры FB2
            # Создаем корневой элемент с корректным объявлением пространств имен
            nsmap = {
            None: "http://www.gribuser.ru/xml/fictionbook/2.0",  # Основное пространство имен (по умолчанию)
            "l": "http://www.w3.org/1999/xlink"                   # Пространство имен для XLink с префиксом 'l'
                    }
            fb2_root = etree.Element("FictionBook", nsmap=nsmap)
            
            # Описание книги
            self._add_description(fb2_root, docx_path)
            
            # Тело книги
            body = etree.SubElement(fb2_root, "body")
            title_element = etree.SubElement(body, "title")
            
            # Обработка содержания
            section = etree.SubElement(body, "section")
            self._process_document(doc, section, settings)
            
            # Обработка изображений
            if settings.convert_images:
                self._extract_images(docx_path, fb2_root)
            
            # Генерация XML
            xml_content = etree.tostring(fb2_root,
                encoding='UTF-8',
                pretty_print=True,
                xml_declaration=True).decode('utf-8')
            
            # Валидация если требуется
            if settings.validate_fb2:
                self._validate_fb2(xml_content)
            
            return xml_content
            
        except Exception as e:
            raise Exception(f"Ошибка конвертации: {str(e)}")
    
    def _add_description(self, root, docx_path):
        """Добавить описание книги"""
        description = etree.SubElement(root, "description")
        title_info = etree.SubElement(description, "title-info")
        
        # Название книги
        book_title = Path(docx_path).stem
        title = etree.SubElement(title_info, "book-title")
        title.text = book_title
        
        # Автор
        author = etree.SubElement(title_info, "author")
        first_name = etree.SubElement(author, "first-name")
        first_name.text = "Автор"
        last_name = etree.SubElement(author, "last-name")
        last_name.text = "Неизвестен"
        
        # Дата
        date = etree.SubElement(title_info, "date")
        date.text = datetime.now().strftime("%Y-%m-%d")
        
        # Язык
        lang = etree.SubElement(title_info, "lang")
        lang.text = "ru"
    
    def _process_document(self, doc, section, settings):
        """Обработать документ"""
        for para in doc.paragraphs:
            if not para.text and settings.remove_empty_lines:
                continue
            
            # Определение стиля
            style = para.style.name.lower() if para.style else ''
            
            # Обработка заголовков
            if 'heading' in style:
                level = 1
                if 'heading 2' in style:
                    level = 2
                elif 'heading 3' in style:
                    level = 3
                
                subtitle = etree.SubElement(section, f"subtitle{'' if level == 1 else str(level)}")
                self._add_text_with_formatting(para, subtitle, settings)
            
            # Обычный текст
            else:
                p = etree.SubElement(section, "p")
                self._add_text_with_formatting(para, p, settings)
    
    def _add_text_with_formatting(self, paragraph, parent, settings):
        """Корректно добавляет текст параграфа с форматированием, сохраняя порядок."""
        if not paragraph.runs:
            if paragraph.text:
                parent.text = paragraph.text
            return

        # Основной алгоритм: проходим по runs и строим структуру последовательно
        current_element = parent  # Начинаем с родительского элемента <p>
        
        for run in paragraph.runs:
            if not run.text:
                continue

            # Определяем, нужен ли тег форматирования для этого run
            if run.bold and settings.preserve_formatting:
                # Если предыдущий элемент уже <strong>, добавляем текст в него
                if current_element.tag == 'strong':
                    current_element.text = (current_element.text or '') + run.text
                else:
                    # Создаем новый тег <strong>
                    strong_elem = etree.SubElement(current_element, "strong")
                    strong_elem.text = run.text
                    current_element = strong_elem
                    
            elif run.italic and settings.preserve_formatting:
                # Аналогично для <emphasis>
                if current_element.tag == 'emphasis':
                    current_element.text = (current_element.text or '') + run.text
                else:
                    emphasis_elem = etree.SubElement(current_element, "emphasis")
                    emphasis_elem.text = run.text
                    current_element = emphasis_elem
            else:
                # Обычный текст
                if current_element == parent:
                    # Добавляем прямо в родительский <p>
                    if parent.text is None:
                        parent.text = run.text
                    else:
                        parent.text += run.text
                else:
                    # Добавляем как хвостовой текст к текущему тегу (например, после </strong>)
                    if current_element.tail is None:
                        current_element.tail = run.text
                    else:
                        current_element.tail += run.text
    
    def _extract_images(self, docx_path, fb2_root):
        """Извлечь изображения из DOCX"""
        try:
            with zipfile.ZipFile(docx_path, 'r') as docx_zip:
                # Поиск изображений
                image_files = [f for f in docx_zip.namelist() 
                             if f.startswith('word/media/') and 
                             f.split('.')[-1].lower() in ['jpg', 'jpeg', 'png', 'gif', 'bmp']]
                
                for img_file in image_files:
                    img_data = docx_zip.read(img_file)
                    ext = img_file.split('.')[-1].lower()
                    
                    # Конвертация в JPG если нужно
                    if ext != 'jpg' and ext != 'jpeg':
                        img = Image.open(BytesIO(img_data))
                        if img.mode in ('RGBA', 'LA', 'P'):
                            img = img.convert('RGB')
                        buffer = BytesIO()
                        img.save(buffer, format='JPEG', quality=90)
                        img_data = buffer.getvalue()
                        ext = 'jpg'
                    
                    # Добавление в FB2
                    binary = etree.SubElement(fb2_root, "binary",
                        id=f"image_{self.image_counter}",
                        content_type=f"image/jpeg")
                    binary.text = base64.b64encode(img_data).decode('ascii')
                    self.image_counter += 1
                    
        except Exception as e:
            print(f"Ошибка извлечения изображений: {e}")
    
    def _validate_fb2(self, xml_content: str):
        """Проверить валидность FB2"""
        try:
            parser = etree.XMLParser(dtd_validation=False)
            etree.fromstring(xml_content.encode('utf-8'), parser)
        except etree.XMLSyntaxError as e:
            raise Exception(f"Ошибка валидации FB2: {str(e)}")


class MainWindow(QMainWindow):
    """Главное окно программы"""
    
    def __init__(self):
        super().__init__()
        self.current_file = None
        self.fb2_content = None
        self.settings = AppSettings.load()
        self.converter = DocxToFb2Converter()
        self.init_ui()
        self.setAcceptDrops(True)
    
    def init_ui(self):
        """Инициализация интерфейса"""
        self.setWindowTitle("DOCtoFB2 - Конвертер для Литрес")
        self.setGeometry(100, 100, 1200, 800)
        
        # Центральный виджет
        central_widget = QWidget()
        self.setCentralWidget(central_widget)
        
        # Основной layout
        main_layout = QVBoxLayout(central_widget)
        
        # Панель инструментов
        self.create_toolbar()
        
        # Разделитель с двумя панелями
        splitter = QSplitter(Qt.Horizontal)
        
        # Левая панель - исходный текст (заглушка)
        self.source_text = QTextEdit()
        self.source_text.setPlaceholderText("Исходный DOC/DOCX файл будет отображен здесь после загрузки")
        self.source_text.setReadOnly(False)
        self.source_text.setPlaceholderText("Исходный DOC/DOCX текст. Можно редактировать перед конвертацией.")
        splitter.addWidget(self.source_text)
        
        # Правая панель - результат FB2
        self.result_text = QTextEdit()
        self.result_text.setPlaceholderText("Результат конвертации в FB2 будет отображен здесь")
        self.result_text.textChanged.connect(self.on_fb2_edited)
        splitter.addWidget(self.result_text)
        
        splitter.setSizes([400, 600])
        main_layout.addWidget(splitter)
        
        # Статус бар
        self.status_bar = QStatusBar()
        self.setStatusBar(self.status_bar)
        self.status_bar.showMessage("Готово к работе")
    
    def create_toolbar(self):
        """Создать панель инструментов"""
        toolbar = QToolBar("Главная панель")
        toolbar.setIconSize(QSize(32, 32))
        self.addToolBar(toolbar)
        
        # Загрузка файла
        load_action = QAction("📂 Загрузить", self)
        load_action.triggered.connect(self.load_file)
        load_action.setShortcut("Ctrl+O")
        toolbar.addAction(load_action)
        
        toolbar.addSeparator()
        
        # Конвертация
        convert_action = QAction("🔄 Конвертировать", self)
        convert_action.triggered.connect(self.convert_file)
        convert_action.setShortcut("Ctrl+R")
        toolbar.addAction(convert_action)
        
        toolbar.addSeparator()
        
        # Сохранить
        save_action = QAction("💾 Сохранить", self)
        save_action.triggered.connect(self.save_fb2)
        save_action.setShortcut("Ctrl+S")
        toolbar.addAction(save_action)
        
        # Сохранить как
        save_as_action = QAction("💾 Сохранить как...", self)
        save_as_action.triggered.connect(self.save_fb2_as)
        save_as_action.setShortcut("Ctrl+Shift+S")
        toolbar.addAction(save_as_action)
        
        toolbar.addSeparator()
        
        # Настройки
        settings_action = QAction("⚙ Настройки", self)
        settings_action.triggered.connect(self.open_settings)
        toolbar.addAction(settings_action)

        # Кнопка статистики
        self.stats_button = QPushButton("📊 Статистика", self)
        self.stats_button.clicked.connect(self.show_statistics)
        toolbar.addWidget(self.stats_button)

        # Добавляем разделитель и кнопку "Очистить всё"
        toolbar.addSeparator()
        self.clear_button = QPushButton(" Очистить всё ")  # Создаем кнопку
        self.clear_button.clicked.connect(self.clear_all_widgets)  # Подключаем обработчик
        toolbar.addWidget(self.clear_button)  # Добавляем кнопку на панель
        
        # О программе
        about_action = QAction("ℹ О программе", self)
        about_action.triggered.connect(self.show_about)
        toolbar.addAction(about_action)
    
    def dragEnterEvent(self, event: QDragEnterEvent):
        """Обработка перетаскивания файла"""
        if event.mimeData().hasUrls():
            urls = event.mimeData().urls()
            if len(urls) == 1:
                file_path = urls[0].toLocalFile()
                if file_path.lower().endswith(('.doc', '.docx')):
                    event.acceptProposedAction()
    
    def dropEvent(self, event: QDropEvent):
        """Обработка сброса файла"""
        urls = event.mimeData().urls()
        if urls:
            file_path = urls[0].toLocalFile()
            if file_path.lower().endswith(('.doc', '.docx')):
                self.load_docx_file(file_path)
    
    def load_file(self):
        """Загрузить DOC/DOCX файл"""
        file_path, _ = QFileDialog.getOpenFileName(
            self,
            "Выберите DOC/DOCX файл",
            "",
            "Документы Word (*.doc *.docx);;Все файлы (*.*)"
        )
        
        if file_path:
            self.load_docx_file(file_path)
    
    def load_docx_file(self, file_path: str):
        """Загрузить и отобразить DOCX файл"""
        try:
            self.current_file = file_path
            
            # Чтение DOCX
            doc = Document(file_path)
            
            # Отображение текста
            text_content = []
            for para in doc.paragraphs:
                if para.text.strip():
                    # Определение стиля
                    style = para.style.name if para.style else 'Обычный'
                    
                    # Отметка заголовков
                    if 'Heading' in style:
                        text_content.append(f"[{style.upper()}] {para.text}")
                    else:
                        text_content.append(para.text)
            
            self.source_text.setPlainText('\n'.join(text_content))
            
            # Статус
            self.status_bar.showMessage(f"Загружен: {os.path.basename(file_path)}")
            
        except Exception as e:
            QMessageBox.critical(self, "Ошибка", f"Не удалось загрузить файл: {str(e)}")
    
    def convert_file(self):
        """Конвертировать файл в FB2"""
        if not self.current_file:
            QMessageBox.warning(self, "Внимание", "Сначала загрузите DOC/DOCX файл")
            return
        
        try:
            # Конвертация
            self.status_bar.showMessage("Конвертация...")
            QApplication.processEvents()
            
            self.fb2_content = self.converter.convert(self.current_file, self.settings)
            
            # Отображение результата
            self.result_text.setPlainText(self.fb2_content)
            
            # Статус
            self.status_bar.showMessage("Конвертация завершена успешно!")
            
            # Автоматическое открытие если настроено
            if self.settings.auto_open_result:
                self.save_fb2_as()
            
        except Exception as e:
            QMessageBox.critical(self, "Ошибка конвертации", str(e))
            self.status_bar.showMessage("Ошибка конвертации")
    
    def save_fb2(self):
        """Сохранить FB2 файл"""
        if not self.fb2_content:
            QMessageBox.warning(self, "Внимание", "Нет данных для сохранения")
            return
        
        # Определение пути сохранения
        if self.settings.default_save_path:
            save_dir = self.settings.default_save_path
        else:
            save_dir = os.path.dirname(self.current_file) if self.current_file else ""
        
        # Имя файла
        if self.current_file:
            base_name = os.path.splitext(os.path.basename(self.current_file))[0]
            default_name = f"{base_name}.fb2"
        else:
            default_name = "книга.fb2"
        
        save_path = os.path.join(save_dir, default_name)
        
        # Запрос подтверждения если файл существует
        if os.path.exists(save_path):
            reply = QMessageBox.question(
                self, "Подтверждение",
                f"Файл {default_name} уже существует. Перезаписать?",
                QMessageBox.Yes | QMessageBox.No
            )
            if reply == QMessageBox.No:
                self.save_fb2_as()
                return
        
        self._save_to_file(save_path)
    
    def save_fb2_as(self):
        """Сохранить FB2 файл как..."""
        if not self.fb2_content:
            QMessageBox.warning(self, "Внимание", "Нет данных для сохранения")
            return
        
        # Определение начального пути
        if self.current_file:
            base_name = os.path.splitext(os.path.basename(self.current_file))[0]
            default_name = f"{base_name}.fb2"
            start_dir = os.path.dirname(self.current_file)
        else:
            default_name = "книга.fb2"
            start_dir = self.settings.default_save_path or ""
        
        # Диалог сохранения
        save_path, _ = QFileDialog.getSaveFileName(
            self,
            "Сохранить FB2 файл",
            os.path.join(start_dir, default_name),
            "FB2 файлы (*.fb2);;Все файлы (*.*)"
        )
        
        if save_path:
            self._save_to_file(save_path)
    
    def _save_to_file(self, file_path: str):
        """Сохранить содержимое в файл"""
        try:
            # Получаем текущее содержимое редактора
            content = self.result_text.toPlainText()
            
            with open(file_path, 'w', encoding='utf-8') as f:
                f.write(content)
            
            self.status_bar.showMessage(f"Файл сохранен: {os.path.basename(file_path)}")
            QMessageBox.information(self, "Сохранено", f"Файл успешно сохранен:\n{file_path}")
            
        except Exception as e:
            QMessageBox.critical(self, "Ошибка", f"Не удалось сохранить файл: {str(e)}")
    
    def on_fb2_edited(self):
        """Обработка редактирования FB2"""
        if self.fb2_content:
            current_content = self.result_text.toPlainText()
            if current_content != self.fb2_content:
                self.status_bar.showMessage("Файл отредактирован - не забудьте сохранить!")
    
    def open_settings(self):
        """Открыть диалог настроек"""
        dialog = SettingsDialog(self)
        if dialog.exec_():
            self.settings = AppSettings.load()


    def show_statistics(self):
        """Показывает статистику для исходного и конвертированного текста."""
        stats_text = "=== СТАТИСТИКА ТЕКСТА ===\n\n"
        
        # Статистика для исходного DOCX
        if self.source_text.toPlainText().strip():
            source_stats = self.calculate_statistics(self.source_text.toPlainText())
            stats_text += "ВХОДНОЙ ТЕКСТ (DOCX):\n"
            stats_text += f"Слов: {source_stats['words']}\n"
            stats_text += f"Знаков (без пробелов): {source_stats['chars_no_spaces']}\n"
            stats_text += f"Знаков (с пробелами): {source_stats['chars_with_spaces']}\n"
            stats_text += f"Абзацев: {source_stats['paragraphs']}\n"
            stats_text += f"Строк: {source_stats['lines']}\n\n"
        
        # Статистика для конвертированного FB2
        if self.result_text.toPlainText().strip():
            # Убираем теги FB2 для чистого текста
            fb2_text = self.strip_fb2_tags(self.result_text.toPlainText())
            result_stats = self.calculate_statistics(fb2_text)
            stats_text += "ВЫХОДНОЙ ТЕКСТ (FB2):\n"
            stats_text += f"Слов: {result_stats['words']}\n"
            stats_text += f"Знаков (без пробелов): {result_stats['chars_no_spaces']}\n"
            stats_text += f"Знаков (с пробелами): {result_stats['chars_with_spaces']}\n"
            stats_text += f"Абзацев: {result_stats['paragraphs']}\n"
            stats_text += f"Строк: {result_stats['lines']}\n"
        
        # Показываем статистику в диалоговом окне
        stats_dialog = QDialog(self)
        stats_dialog.setWindowTitle("Статистика текста")
        stats_dialog.setFixedSize(400, 300)
        
        layout = QVBoxLayout()
        text_edit = QTextEdit()
        text_edit.setPlainText(stats_text)
        text_edit.setReadOnly(True)
        layout.addWidget(text_edit)
        
        # Кнопка копирования
        copy_btn = QPushButton("Копировать в буфер")
        copy_btn.clicked.connect(lambda: QApplication.clipboard().setText(stats_text))
        layout.addWidget(copy_btn)
        
        stats_dialog.setLayout(layout)
        stats_dialog.exec_()

    def calculate_statistics(self, text):
        """Вычисляет статистику текста."""
        lines = text.count('\n') + 1
        paragraphs = len([p for p in text.split('\n') if p.strip()])
        words = len(text.split())
        chars_with_spaces = len(text)
        chars_no_spaces = len(text.replace(" ", "").replace("\n", "").replace("\t", ""))
        
        return {
            'lines': lines, 'paragraphs': paragraphs, 'words': words,
            'chars_with_spaces': chars_with_spaces, 'chars_no_spaces': chars_no_spaces
        }

    def strip_fb2_tags(self, fb2_text):
        """Удаляет теги FB2 для подсчета статистики."""
        import re
        # Удаляем XML теги
        clean_text = re.sub(r'<[^>]+>', '', fb2_text)
        # Заменяем XML сущности
        clean_text = clean_text.replace('&lt;', '<').replace('&gt;', '>')
        return clean_text    


    def clear_all_widgets(self):
        """Очищает все текстовые поля и сбрасывает состояние."""
        try:
            # Очищаем основное текстовое поле с исходным текстом
            if hasattr(self, 'source_text'):
                self.source_text.clear()
            else:
                print("[DEBUG] Атрибут source_text не найден")
            
            # Очищаем поле с результатом конвертации (FB2)
            if hasattr(self, 'result_text'):
                self.result_text.clear()
            else:
                print("[DEBUG] Атрибут result_text не найден")
            
            # Очищаем статусную строку
            if hasattr(self, 'status_bar') and self.status_bar:
                self.status_bar.clearMessage()
            
            # Сбрасываем путь к текущему файлу
            self.current_file = None
            
            # Обновляем надпись о загруженном файле (если есть)
            if hasattr(self, 'file_label'):
                self.file_label.setText("Файл не выбран")
            else:
                print("[DEBUG] Атрибут file_label не найден")
            
            # Добавляем запись в лог (если есть)
            if hasattr(self, 'log_text'):
                self.log_text.append("> Все поля очищены.")
            else:
                print("[DEBUG] Атрибут log_text не найден")
                
        except Exception as e:
            # Выводим ошибку в консоль для отладки
            print(f"[ОШИБКА в clear_all_widgets]: {str(e)}")
            # Можно показать сообщение пользователю
            import traceback
            traceback.print_exc()  # Полная трассировка ошибки       
    
    def show_about(self):
        """Показать информацию о программе"""
        about_text = """
        <h2>DOCtoFB2 - Конвертер для Литрес Самиздат</h2>
        <p><b>Автор:</b> VUS HAAR (C)</p>
        <p><b>Версия:</b> 1.1.3</p>
        <p><b>Описание:</b> Программа для конвертации файлов DOC/DOCX в формат FB2 
        с соблюдением правил платформы Литрес Самиздат.</p>
        
        <h3>Основные возможности:</h3>
        <ul>
        <li>Конвертация DOC/DOCX в FB2</li>
        <li>Поддержка стилей заголовков</li>
        <li>Обработка изображений</li>
        <li>Редактирование FB2 файлов</li>
        <li>Drag-and-drop загрузка файлов</li>
        <li>Настройка параметров конвертации</li>
        </ul>
        
        <h3>Правила Литрес:</h3>
        <p>Программа учитывает основные требования Литрес для публикации книг:</p>
        <ul>
        <li>Правильное оформление заголовков</li>
        <li>Обработка изображений</li>
        <li>Удаление лишнего форматирования</li>
        </ul>
        
        <p><i>Для получения наилучших результатов следуйте рекомендациям Литрес 
        при подготовке исходного документа.</i></p>
        """
        
        QMessageBox.about(self, "О программе DOCtoFB2", about_text)


def main():
    """Точка входа в программу"""
    # Настройка приложения
    app = QApplication(sys.argv)
    app.setApplicationName("DOCtoFB2")
    app.setOrganizationName("LitresTools")
    
    # Создание и отображение главного окна
    window = MainWindow()
    window.show()
    
    # Обработка аргументов командной строки
    if len(sys.argv) > 1:
        file_path = sys.argv[1]
        if os.path.exists(file_path) and file_path.lower().endswith(('.doc', '.docx')):
            window.load_docx_file(file_path)
    
    sys.exit(app.exec_())


if __name__ == "__main__":
    main()